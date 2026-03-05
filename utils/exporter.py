# utils/exporter.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


def generate_word_doc(project_name: str, chapters: list) -> BytesIO:
    """
    生成 Word 文档流
    :param project_name: 小说标题
    :param chapters: 章节列表 [{'title':..., 'content':...}, ...]
    :return: BytesIO 对象
    """
    doc = Document()

    # 1. 封面标题
    heading = doc.add_heading(project_name, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 2. 遍历章节
    for chap in chapters:
        # 章节标题 (一级标题)
        title = chap.get('title', f"第 {chap.get('chapter_num')} 章")
        doc.add_heading(title, level=1)

        # 章节正文
        content = chap.get('content', '')
        if content:
            # 简单处理换行
            paragraphs = content.split('\n')
            for p_text in paragraphs:
                if p_text.strip():
                    p = doc.add_paragraph(p_text.strip())
                    p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
                    p.paragraph_format.space_after = Pt(10)  # 段后间距

        doc.add_page_break()

    # 3. 保存到内存流
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def generate_txt_doc(project_name: str, chapters: list) -> BytesIO:
    """生成 TXT 文档流"""
    output = f"《{project_name}》\n\n"

    for chap in chapters:
        title = chap.get('title', '')
        content = chap.get('content', '')
        output += f"=== {title} ===\n\n"
        output += content + "\n\n"
        output += "-" * 20 + "\n\n"

    file_stream = BytesIO(output.encode('utf-8'))
    file_stream.seek(0)
    return file_stream